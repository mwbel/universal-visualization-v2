"""
万物可视化 v2.0 - 页面级内容分析器
对文档进行逐页识别，判断每页内容的学科类型，并匹配相应的可视化模板
"""

import os
import re
from typing import Dict, List, Optional, Tuple, Any
from dataclasses import dataclass
from datetime import datetime

# PDF处理库
try:
    import PyPDF2
    import pdfplumber
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

# 文档处理库
try:
    import docx
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

# 图片处理库（用于扫描文档的OCR）
try:
    import cv2
    import numpy as np
    from PIL import Image
    import pytesseract
    OCR_SUPPORT = True
except ImportError:
    OCR_SUPPORT = False

from data.subject_keywords_database import SubjectKeywordDatabase, SubjectType, keyword_database
from agents.file_analysis_agent import FileAnalysisResult

@dataclass
class PageContent:
    """页面内容数据结构"""
    page_number: int
    text_content: str
    image_paths: List[str]  # 页面中的图片路径
    tables: List[List[List[str]]]  # 表格数据
    metadata: Dict[str, Any]  # 页面元数据
    word_count: int
    character_count: int

@dataclass
class PageAnalysisResult:
    """页面分析结果"""
    page_number: int
    identified_subjects: List[Tuple[SubjectType, float, List[str]]]  # 学科类型、置信度、匹配关键词
    primary_subject: Optional[SubjectType]  # 主要学科
    confidence_score: float  # 主要学科的置信度
    recommended_templates: List[str]  # 推荐的可视化模板
    key_concepts: List[str]  # 关键概念
    content_type: str  # 内容类型：text, mixed, formula, diagram, table等
    difficulty_level: str  # 难度级别：basic, intermediate, advanced
    has_visualizable_content: bool  # 是否包含可可视化内容

class PageLevelAnalyzer:
    """页面级内容分析器"""

    def __init__(self, keyword_db: Optional[SubjectKeywordDatabase] = None):
        self.keyword_db = keyword_db or keyword_database
        self.supported_formats = ['.pdf', '.docx', '.txt', '.md']

    def analyze_document_pages(self, file_path: str, file_id: str) -> Dict[str, Any]:
        """
        分析文档的所有页面

        Args:
            file_path: 文件路径
            file_id: 文件ID

        Returns:
            Dict: 包含所有页面分析结果的字典
        """
        file_ext = os.path.splitext(file_path)[1].lower()

        try:
            if file_ext == '.pdf':
                return self._analyze_pdf_pages(file_path, file_id)
            elif file_ext in ['.txt', '.md']:
                return self._analyze_text_pages(file_path, file_id)
            elif file_ext == '.docx':
                return self._analyze_docx_pages(file_path, file_id)
            else:
                return {"error": f"不支持的文件格式: {file_ext}"}

        except Exception as e:
            return {"error": f"页面分析失败: {str(e)}"}

    def _analyze_pdf_pages(self, file_path: str, file_id: str) -> Dict[str, Any]:
        """分析PDF文档的每一页"""
        if not PDF_SUPPORT:
            return {"error": "PDF分析需要安装PyPDF2和pdfplumber库"}

        pages_content = []
        pages_analysis = []

        try:
            # 优先使用pdfplumber，它能更好地提取表格和图片信息
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    # 提取文本
                    text = page.extract_text() or ""

                    # 提取表格
                    tables = page.extract_tables() or []

                    # 提取图片信息
                    images = []
                    if hasattr(page, 'images'):
                        images = [img.get('filename', f'image_{len(images)}') for img in page.images]

                    # 创建页面内容对象
                    page_content = PageContent(
                        page_number=page_num,
                        text_content=text,
                        image_paths=images,
                        tables=tables,
                        metadata={
                            'width': page.width,
                            'height': page.height,
                            'bbox': page.bbox
                        },
                        word_count=len(text.split()),
                        character_count=len(text)
                    )
                    pages_content.append(page_content)

                    # 分析页面内容
                    analysis = self._analyze_page_content(page_content)
                    pages_analysis.append(analysis)

        except ImportError:
            # 如果没有pdfplumber，使用PyPDF2作为后备
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    text = page.extract_text() or ""

                    page_content = PageContent(
                        page_number=page_num,
                        text_content=text,
                        image_paths=[],
                        tables=[],
                        metadata={'page_number': page_num},
                        word_count=len(text.split()),
                        character_count=len(text)
                    )
                    pages_content.append(page_content)

                    analysis = self._analyze_page_content(page_content)
                    pages_analysis.append(analysis)

        # 生成文档级别的总结
        document_summary = self._generate_document_summary(pages_analysis, pages_content)

        return {
            "file_id": file_id,
            "total_pages": len(pages_content),
            "pages_content": [self._page_content_to_dict(pc) for pc in pages_content],
            "pages_analysis": [self._page_analysis_to_dict(pa) for pa in pages_analysis],
            "document_summary": document_summary,
            "analysis_time": datetime.now().isoformat()
        }

    def _analyze_text_pages(self, file_path: str, file_id: str) -> Dict[str, Any]:
        """分析文本文档的页面（按段落或章节分割）"""
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()

        # 按空行分割页面/段落
        sections = re.split(r'\n\s*\n', content)
        pages_content = []
        pages_analysis = []

        for page_num, section in enumerate(sections, 1):
            if section.strip():  # 跳过空段落
                page_content = PageContent(
                    page_number=page_num,
                    text_content=section.strip(),
                    image_paths=[],
                    tables=[],
                    metadata={'section_number': page_num},
                    word_count=len(section.split()),
                    character_count=len(section)
                )
                pages_content.append(page_content)

                analysis = self._analyze_page_content(page_content)
                pages_analysis.append(analysis)

        document_summary = self._generate_document_summary(pages_analysis, pages_content)

        return {
            "file_id": file_id,
            "total_pages": len(pages_content),
            "pages_content": [self._page_content_to_dict(pc) for pc in pages_content],
            "pages_analysis": [self._page_analysis_to_dict(pa) for pa in pages_analysis],
            "document_summary": document_summary,
            "analysis_time": datetime.now().isoformat()
        }

    def _analyze_docx_pages(self, file_path: str, file_id: str) -> Dict[str, Any]:
        """分析Word文档的页面"""
        if not DOCX_SUPPORT:
            return {"error": "Word文档分析需要安装python-docx库"}

        doc = docx.Document(file_path)
        pages_content = []
        pages_analysis = []

        # 将段落按合理的方式分组为页面
        paragraphs_per_page = 20  # 估计每页段落数
        current_page_text = []
        page_num = 1

        for paragraph in doc.paragraphs:
            current_page_text.append(paragraph.text)

            # 每20个段落或遇到分页符时创建新页面
            if len(current_page_text) >= paragraphs_per_page or 'PAGE_BREAK' in paragraph.text:
                page_text = '\n'.join(current_page_text)

                page_content = PageContent(
                    page_number=page_num,
                    text_content=page_text,
                    image_paths=[],
                    tables=[],
                    metadata={'paragraph_count': len(current_page_text)},
                    word_count=len(page_text.split()),
                    character_count=len(page_text)
                )
                pages_content.append(page_content)

                analysis = self._analyze_page_content(page_content)
                pages_analysis.append(analysis)

                current_page_text = []
                page_num += 1

        # 处理最后剩余的段落
        if current_page_text:
            page_text = '\n'.join(current_page_text)
            page_content = PageContent(
                page_number=page_num,
                text_content=page_text,
                image_paths=[],
                tables=[],
                metadata={'paragraph_count': len(current_page_text)},
                word_count=len(page_text.split()),
                character_count=len(page_text)
            )
            pages_content.append(page_content)

            analysis = self._analyze_page_content(page_content)
            pages_analysis.append(analysis)

        document_summary = self._generate_document_summary(pages_analysis, pages_content)

        return {
            "file_id": file_id,
            "total_pages": len(pages_content),
            "pages_content": [self._page_content_to_dict(pc) for pc in pages_content],
            "pages_analysis": [self._page_analysis_to_dict(pa) for pa in pages_analysis],
            "document_summary": document_summary,
            "analysis_time": datetime.now().isoformat()
        }

    def _analyze_page_content(self, page_content: PageContent) -> PageAnalysisResult:
        """分析单个页面的内容"""
        text = page_content.text_content

        # 识别学科类型
        identified_subjects = self.keyword_db.identify_subjects(text, min_confidence=0.1)

        # 确定主要学科
        primary_subject = None
        confidence_score = 0.0
        if identified_subjects:
            primary_subject, confidence_score, matched_keywords = identified_subjects[0]

        # 获取推荐的模板
        recommended_templates = []
        if primary_subject:
            recommended_templates = self.keyword_db.get_subject_templates(primary_subject)

        # 提取关键概念
        key_concepts = self._extract_key_concepts(text, identified_subjects)

        # 判断内容类型
        content_type = self._detect_content_type(text, page_content)

        # 评估难度级别
        difficulty_level = self._assess_difficulty_level(text, key_concepts)

        # 判断是否包含可可视化内容
        has_visualizable_content = self._has_visualizable_content(
            text, page_content.tables, page_content.image_paths
        )

        return PageAnalysisResult(
            page_number=page_content.page_number,
            identified_subjects=identified_subjects,
            primary_subject=primary_subject,
            confidence_score=confidence_score,
            recommended_templates=recommended_templates,
            key_concepts=key_concepts,
            content_type=content_type,
            difficulty_level=difficulty_level,
            has_visualizable_content=has_visualizable_content
        )

    def _extract_key_concepts(self, text: str, identified_subjects: List[Tuple[SubjectType, float, List[str]]]) -> List[str]:
        """提取关键概念"""
        concepts = []

        # 从识别的学科中获取关键词
        for subject_type, confidence, keywords in identified_subjects:
            if confidence > 0.3:  # 只考虑置信度较高的学科
                concepts.extend(keywords[:5])  # 每个学科取前5个关键词

        # 使用简单的术语提取
        # 查找大写词汇、专业术语等
        technical_terms = re.findall(r'\b[A-Z][a-zA-Z]+\b', text)
        concepts.extend(technical_terms)

        # 查找中文专业术语（2-6个字符的中文词组）
        chinese_terms = re.findall(r'[\u4e00-\u9fff]{2,6}', text)
        # 过滤常见词汇，保留可能的专业术语
        filtered_chinese = [term for term in chinese_terms if len(term) >= 3]
        concepts.extend(filtered_chinese[:10])

        # 去重并限制数量
        unique_concepts = list(set(concepts))
        return unique_concepts[:20]

    def _detect_content_type(self, text: str, page_content: PageContent) -> str:
        """检测内容类型"""
        # 检测数学公式
        if re.search(r'\\[a-zA-Z]+\{[^}]*\}|\\frac\{[^}]*\}\{[^}]*\}|∑∏∫√±≤≥≠', text):
            return "formula_heavy"

        # 检测代码
        if re.search(r'\b(function|def|class|if|else|for|while|return)\b|\{|\}|;', text):
            return "code"

        # 检测表格
        if page_content.tables:
            return "table_heavy"

        # 检测图表描述
        if re.search(r'(图表|图|表|图像|示意图|流程图)', text, re.IGNORECASE):
            return "diagram"

        # 检测列表
        if re.search(r'^\s*[-*+]\s+', text, re.MULTILINE):
            return "list"

        # 检测长文本
        if len(text.split()) > 100:
            return "text_heavy"

        return "mixed"

    def _assess_difficulty_level(self, text: str, key_concepts: List[str]) -> str:
        """评估难度级别"""
        difficulty_score = 0

        # 基于关键词复杂度
        complex_terms = [
            '微分', '积分', '矩阵', '特征值', '量子', '相对论', '催化', '新陈代谢',
            '算法', '复杂度', '递归', '多态', '封装', '继承', '抽象'
        ]

        for concept in key_concepts:
            for complex_term in complex_terms:
                if complex_term in concept:
                    difficulty_score += 1

        # 基于句子长度和复杂度
        sentences = re.split(r'[。.!?！？]', text)
        avg_sentence_length = sum(len(s) for s in sentences) / max(len(sentences), 1)

        if avg_sentence_length > 50:
            difficulty_score += 2
        elif avg_sentence_length > 30:
            difficulty_score += 1

        # 基于专业术语密度
        technical_density = len(key_concepts) / max(len(text.split()), 1) * 100
        if technical_density > 20:
            difficulty_score += 2
        elif technical_density > 10:
            difficulty_score += 1

        # 确定难度级别
        if difficulty_score >= 5:
            return "advanced"
        elif difficulty_score >= 3:
            return "intermediate"
        else:
            return "basic"

    def _has_visualizable_content(self, text: str, tables: List, images: List[str]) -> bool:
        """判断页面是否包含可可视化内容"""
        # 有数据表格
        if tables:
            return True

        # 有图片
        if images:
            return True

        # 包含数值数据
        numbers = re.findall(r'\b\d+\.?\d*\b', text)
        if len(numbers) >= 3:
            return True

        # 包含百分比
        if re.search(r'\b\d+%\b', text):
            return True

        # 包含时间序列数据
        if re.search(r'\b(19|20)\d{2}\b', text):  # 年份
            return True

        # 包含对比性词汇
        comparison_words = ['比较', '对比', '增加', '减少', '上升', '下降', '增长', '变化']
        if any(word in text for word in comparison_words):
            return True

        return False

    def _generate_document_summary(self, pages_analysis: List[PageAnalysisResult], pages_content: List[PageContent]) -> Dict[str, Any]:
        """生成文档级别的总结"""
        total_pages = len(pages_analysis)

        # 统计学科分布
        subject_distribution = {}
        for analysis in pages_analysis:
            if analysis.primary_subject:
                subject_name = analysis.primary_subject.value
                if subject_name not in subject_distribution:
                    subject_distribution[subject_name] = {"count": 0, "total_confidence": 0.0}
                subject_distribution[subject_name]["count"] += 1
                subject_distribution[subject_name]["total_confidence"] += analysis.confidence_score

        # 计算平均置信度
        for subject in subject_distribution:
            count = subject_distribution[subject]["count"]
            total_conf = subject_distribution[subject]["total_confidence"]
            subject_distribution[subject]["average_confidence"] = total_conf / count

        # 统计内容类型分布
        content_type_distribution = {}
        for analysis in pages_analysis:
            content_type = analysis.content_type
            if content_type not in content_type_distribution:
                content_type_distribution[content_type] = 0
            content_type_distribution[content_type] += 1

        # 统计难度级别分布
        difficulty_distribution = {"basic": 0, "intermediate": 0, "advanced": 0}
        for analysis in pages_analysis:
            difficulty_distribution[analysis.difficulty_level] += 1

        # 统计可可视化页面
        visualizable_pages = sum(1 for analysis in pages_analysis if analysis.has_visualizable_content)

        # 收集所有关键概念
        all_concepts = []
        for analysis in pages_analysis:
            all_concepts.extend(analysis.key_concepts)

        # 统计最频繁的概念
        from collections import Counter
        concept_counter = Counter(all_concepts)
        top_concepts = concept_counter.most_common(20)

        return {
            "total_pages": total_pages,
            "subject_distribution": subject_distribution,
            "content_type_distribution": content_type_distribution,
            "difficulty_distribution": difficulty_distribution,
            "visualizable_pages": visualizable_pages,
            "visualizable_ratio": visualizable_pages / total_pages if total_pages > 0 else 0,
            "top_concepts": [{"concept": concept, "frequency": freq} for concept, freq in top_concepts],
            "primary_subjects": sorted(
                [(subj, data["count"]) for subj, data in subject_distribution.items()],
                key=lambda x: x[1], reverse=True
            )[:3]
        }

    def _page_content_to_dict(self, page_content: PageContent) -> Dict[str, Any]:
        """将PageContent对象转换为字典"""
        return {
            "page_number": page_content.page_number,
            "text_content": page_content.text_content,
            "image_paths": page_content.image_paths,
            "tables": page_content.tables,
            "metadata": page_content.metadata,
            "word_count": page_content.word_count,
            "character_count": page_content.character_count
        }

    def _page_analysis_to_dict(self, analysis: PageAnalysisResult) -> Dict[str, Any]:
        """将PageAnalysisResult对象转换为字典"""
        return {
            "page_number": analysis.page_number,
            "identified_subjects": [
                {
                    "subject_type": subject.value,
                    "confidence": confidence,
                    "matched_keywords": keywords
                }
                for subject, confidence, keywords in analysis.identified_subjects
            ],
            "primary_subject": analysis.primary_subject.value if analysis.primary_subject else None,
            "confidence_score": analysis.confidence_score,
            "recommended_templates": analysis.recommended_templates,
            "key_concepts": analysis.key_concepts,
            "content_type": analysis.content_type,
            "difficulty_level": analysis.difficulty_level,
            "has_visualizable_content": analysis.has_visualizable_content
        }

# 全局页面分析器实例
page_analyzer = PageLevelAnalyzer()